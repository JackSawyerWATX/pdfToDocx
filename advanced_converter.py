import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re

def extract_with_tables_and_formatting(pdf_path):
    """
    Advanced PDF to DOCX converter that preserves:
    - Tables
    - Text positioning
    - Font size hints
    - Bullet points
    - Better paragraph structure
    """
    pdf = pdfplumber.open(pdf_path)
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    for page_num, page in enumerate(pdf.pages, 1):
        if page_num > 1:
            doc.add_page_break()
        
        print(f"Processing page {page_num}...")
        
        # Try to extract tables first
        tables = page.extract_tables()
        if tables:
            for table_data in tables:
                if table_data and len(table_data) > 0:
                    # Create table in Word
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            if cell_text:
                                table.rows[i].cells[j].text = str(cell_text).strip()
                    
                    doc.add_paragraph()  # Add space after table
        
        # Extract text with character-level details for better formatting
        chars = page.chars
        if chars:
            # Group characters by line
            lines = {}
            for char in chars:
                y = round(char['y0'], 1)  # Round to group chars on same line
                if y not in lines:
                    lines[y] = []
                lines[y].append(char)
            
            # Sort lines by vertical position (top to bottom)
            sorted_lines = sorted(lines.items(), key=lambda x: x[0], reverse=True)
            
            current_paragraph_text = []
            
            for y_pos, line_chars in sorted_lines:
                # Sort characters in line by horizontal position
                line_chars.sort(key=lambda c: c['x0'])
                
                # Extract text from line
                line_text = ''.join(char['text'] for char in line_chars).strip()
                
                if not line_text:
                    # Empty line - finish current paragraph
                    if current_paragraph_text:
                        para_text = ' '.join(current_paragraph_text)
                        add_formatted_paragraph(doc, para_text, line_chars)
                        current_paragraph_text = []
                    continue
                
                # Detect if this should be a new paragraph (large font, all caps, etc.)
                avg_font_size = sum(char.get('size', 12) for char in line_chars) / len(line_chars)
                is_likely_header = (
                    avg_font_size > 14 or 
                    line_text.isupper() or 
                    len(line_text) < 50
                )
                
                if is_likely_header and current_paragraph_text:
                    # Finish previous paragraph
                    para_text = ' '.join(current_paragraph_text)
                    add_formatted_paragraph(doc, para_text, [])
                    current_paragraph_text = []
                    
                    # Add header
                    if line_text.isupper() and len(line_text) < 30:
                        heading = doc.add_heading(line_text, level=2)
                    else:
                        para = doc.add_paragraph()
                        run = para.add_run(line_text)
                        if avg_font_size > 14:
                            run.font.size = Pt(int(avg_font_size))
                            run.bold = True
                else:
                    current_paragraph_text.append(line_text)
            
            # Add any remaining paragraph
            if current_paragraph_text:
                para_text = ' '.join(current_paragraph_text)
                add_formatted_paragraph(doc, para_text, [])
        
        else:
            # Fallback to basic text extraction
            text = page.extract_text()
            if text:
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        if detect_bullet_point(line):
                            add_bullet_point(doc, line)
                        elif line.isupper() and len(line) < 50:
                            doc.add_heading(line, level=2)
                        else:
                            doc.add_paragraph(line)
    
    pdf.close()
    return doc

def add_formatted_paragraph(doc, text, chars):
    """Add a paragraph with basic formatting detection"""
    text = text.strip()
    if not text:
        return
    
    if detect_bullet_point(text):
        add_bullet_point(doc, text)
    elif text.isupper() and len(text) < 50:
        doc.add_heading(text, level=2)
    else:
        para = doc.add_paragraph(text)
        # Try to detect and preserve some formatting
        if any(keyword in text.lower() for keyword in ['summary', 'objective', 'experience', 'education', 'skills']):
            run = para.runs[0] if para.runs else para.add_run()
            run.bold = True

def detect_bullet_point(text):
    """Detect if text should be a bullet point"""
    bullet_indicators = ['•', '◦', '-', '*', '▪', '▫']
    return any(text.startswith(indicator) for indicator in bullet_indicators) or \
           re.match(r'^\s*[\d]+\.', text) or \
           re.match(r'^\s*[a-zA-Z]\.', text)

def add_bullet_point(doc, text):
    """Add text as a bullet point"""
    # Remove common bullet characters
    text = re.sub(r'^[\s•◦\-*▪▫]+', '', text).strip()
    text = re.sub(r'^\d+\.\s*', '', text).strip()  # Remove numbered bullets
    text = re.sub(r'^[a-zA-Z]\.\s*', '', text).strip()  # Remove lettered bullets
    
    if text:
        para = doc.add_paragraph(text, style='List Bullet')

# Main execution
if __name__ == "__main__":
    pdf_path = input("Enter the path to the PDF file: ").strip().strip('"').strip("'")
    
    try:
        print("Starting advanced conversion with formatting preservation...")
        doc = extract_with_tables_and_formatting(pdf_path)
        
        output_filename = "PDF-Advanced-Formatted.docx"
        doc.save(output_filename)
        print(f"✅ Advanced conversion complete! Saved as '{output_filename}'")
        print("This version preserves:")
        print("- Tables and their structure")
        print("- Headers and text hierarchy")
        print("- Bullet points and lists")
        print("- Basic font size differences")
        print("- Paragraph spacing")
        
    except Exception as e:
        print(f"❌ Error during advanced conversion: {e}")
        print("Please try the basic converter (pdfDocx.py) instead.")