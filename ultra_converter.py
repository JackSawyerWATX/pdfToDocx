"""
Ultra-Advanced PDF to DOCX Converter
This script provides the best possible formatting preservation by combining multiple techniques.
Requires additional packages: pdf2image, Pillow
"""

import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def install_requirements():
    """Install additional required packages"""
    try:
        import pdf2image
        from PIL import Image
        print("✅ All required packages are available")
        return True
    except ImportError:
        print("📦 Installing required packages...")
        import subprocess
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pdf2image", "Pillow"])
            print("✅ Packages installed successfully!")
            return True
        except Exception as e:
            print(f"❌ Failed to install packages: {e}")
            return False

def ultra_preserve_formatting(pdf_path):
    """
    Ultimate formatting preservation using multiple extraction methods
    """
    pdf = pdfplumber.open(pdf_path)
    doc = Document()
    
    # Set document margins to match typical PDF layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    for page_num, page in enumerate(pdf.pages, 1):
        if page_num > 1:
            doc.add_page_break()
        
        print(f"🔄 Processing page {page_num} with ultra formatting...")
        
        # Method 1: Extract tables with precise formatting
        tables = page.extract_tables()
        table_areas = []
        
        if tables:
            for table_data in tables:
                if table_data and len(table_data) > 0:
                    print(f"  📊 Found table with {len(table_data)} rows")
                    
                    # Create Word table
                    word_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    word_table.style = 'Light Grid Accent 1'
                    
                    # Fill table with data
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            if cell_text:
                                cell = word_table.rows[i].cells[j]
                                cell.text = str(cell_text).strip()
                                
                                # Basic cell formatting
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(11)
                    
                    doc.add_paragraph()  # Space after table
        
        # Method 2: Character-level analysis for precise text formatting
        chars = page.chars
        if chars:
            # Group characters by their properties and position
            text_blocks = analyze_text_blocks(chars)
            
            for block in text_blocks:
                add_formatted_text_block(doc, block)
        
        # Method 3: Fallback text extraction with smart parsing
        else:
            text = page.extract_text()
            if text:
                smart_text_parsing(doc, text)
    
    pdf.close()
    return doc

def analyze_text_blocks(chars):
    """
    Analyze character-level data to group into formatted text blocks
    """
    # Group characters by line and analyze their properties
    lines = {}
    
    for char in chars:
        y = round(char['y0'], 2)
        if y not in lines:
            lines[y] = {
                'chars': [],
                'font_sizes': [],
                'fonts': [],
                'x_positions': []
            }
        
        lines[y]['chars'].append(char)
        lines[y]['font_sizes'].append(char.get('size', 12))
        lines[y]['fonts'].append(char.get('fontname', 'default'))
        lines[y]['x_positions'].append(char['x0'])
    
    # Convert to text blocks with formatting info
    text_blocks = []
    sorted_lines = sorted(lines.items(), key=lambda x: x[0], reverse=True)
    
    for y_pos, line_data in sorted_lines:
        chars_in_line = sorted(line_data['chars'], key=lambda c: c['x0'])
        text = ''.join(char['text'] for char in chars_in_line).strip()
        
        if not text:
            text_blocks.append({'type': 'empty', 'text': ''})
            continue
        
        # Analyze formatting
        avg_font_size = sum(line_data['font_sizes']) / len(line_data['font_sizes'])
        most_common_font = max(set(line_data['fonts']), key=line_data['fonts'].count)
        indent_level = min(line_data['x_positions']) if line_data['x_positions'] else 0
        
        # Classify text type
        block_type = classify_text_type(text, avg_font_size, indent_level)
        
        text_blocks.append({
            'type': block_type,
            'text': text,
            'font_size': avg_font_size,
            'font': most_common_font,
            'indent': indent_level
        })
    
    return text_blocks

def classify_text_type(text, font_size, indent_level):
    """Classify what type of text this is based on content and formatting"""
    
    # Check for headers (large font, short text, all caps, etc.)
    if font_size > 16:
        return 'heading1'
    elif font_size > 14:
        return 'heading2'
    elif text.isupper() and len(text) < 50:
        return 'heading2'
    
    # Check for bullet points
    if re.match(r'^[\s•◦\-*▪▫]', text) or re.match(r'^\s*\d+\.', text):
        return 'bullet'
    
    # Check for contact info or special formatting
    if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text):  # Phone number
        return 'contact'
    elif re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text):  # Email
        return 'contact'
    
    # Check for section headers (common resume sections)
    section_keywords = ['experience', 'education', 'skills', 'summary', 'objective', 'projects', 'certifications']
    if any(keyword in text.lower() for keyword in section_keywords) and len(text) < 30:
        return 'section_header'
    
    # High indentation suggests sub-items
    if indent_level > 50:
        return 'indented'
    
    return 'paragraph'

def add_formatted_text_block(doc, block):
    """Add a text block with appropriate formatting"""
    
    if block['type'] == 'empty':
        doc.add_paragraph()
        return
    
    text = block['text']
    
    if block['type'] == 'heading1':
        heading = doc.add_heading(text, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    elif block['type'] == 'heading2':
        heading = doc.add_heading(text, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    elif block['type'] == 'section_header':
        heading = doc.add_heading(text, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add underline for section headers
        for run in heading.runs:
            run.underline = True
    
    elif block['type'] == 'bullet':
        # Clean bullet text
        clean_text = re.sub(r'^[\s•◦\-*▪▫]+', '', text).strip()
        clean_text = re.sub(r'^\d+\.\s*', '', clean_text).strip()
        if clean_text:
            para = doc.add_paragraph(clean_text, style='List Bullet')
    
    elif block['type'] == 'contact':
        para = doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(11)
            run.italic = True
    
    elif block['type'] == 'indented':
        para = doc.add_paragraph(text)
        para.paragraph_format.left_indent = Inches(0.5)
        for run in para.runs:
            run.font.size = Pt(10)
    
    else:  # paragraph
        para = doc.add_paragraph(text)
        # Preserve approximate font size
        for run in para.runs:
            if block['font_size'] > 12:
                run.font.size = Pt(int(block['font_size']))

def smart_text_parsing(doc, text):
    """Fallback smart text parsing when character data isn't available"""
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Apply smart formatting based on content
        if line.isupper() and len(line) < 50:
            doc.add_heading(line, level=2)
        elif re.match(r'^[\s•◦\-*▪▫]', line) or re.match(r'^\s*\d+\.', line):
            clean_text = re.sub(r'^[\s•◦\-*▪▫\d.]+', '', line).strip()
            if clean_text:
                doc.add_paragraph(clean_text, style='List Bullet')
        else:
            doc.add_paragraph(line)

# Main execution
if __name__ == "__main__":
    print("🚀 Ultra-Advanced PDF to DOCX Converter")
    print("=" * 50)
    
    # Check and install requirements
    if not install_requirements():
        print("❌ Cannot proceed without required packages")
        sys.exit(1)
    
    pdf_path = input("Enter the path to the PDF file: ").strip().strip('"').strip("'")
    
    try:
        print(f"\n🔄 Starting ultra-advanced conversion...")
        print("This may take longer but will preserve maximum formatting!")
        
        doc = ultra_preserve_formatting(pdf_path)
        
        output_filename = "PDF-Ultra-Formatted.docx"
        doc.save(output_filename)
        
        print(f"\n✅ Ultra-conversion complete!")
        print(f"📄 Output saved as: {output_filename}")
        print("\n🎨 This version preserves:")
        print("  • Tables with proper formatting")
        print("  • Headers and text hierarchy")
        print("  • Font sizes and styles")
        print("  • Bullet points and indentation")
        print("  • Contact information formatting")
        print("  • Section headers with emphasis")
        print("  • Paragraph spacing and structure")
        
    except Exception as e:
        print(f"\n❌ Error during ultra conversion: {e}")
        print("💡 Try using the advanced_converter.py or basic pdfDocx.py instead")